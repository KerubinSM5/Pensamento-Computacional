from docx import Document

numero = int(input('Tabuada de qual número? '))

doc = Document()

doc.add_heading('Tabuada de {}'.format(numero), 0)

tab = doc.add_table(rows=1, cols=5)
tab.style="Medium List 1 Accent 2"
cels = tab.rows[0].cells
cels[0].text = 'Núm'
cels[1].text = "Opr"
cels[2].text = "Rep"
cels[3].text = "="
cels[4].text = "Res"

for rep in range(11):
   dados = tab.add_row().cells
   dados[0].text = str(numero)
   dados[1].text = 'x'
   dados[2].text = str(rep)
   dados[3].text = '='
   dados[4].text = str(numero * rep)

doc.save('tabuada{}.docx'.format(numero))
