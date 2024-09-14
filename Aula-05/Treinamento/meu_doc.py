from docx import Document
from docx.shared import Inches

doc = Document()

doc.add_heading('Estes são os seus dados', 0)

nome = input("Qual o seu nome? ")
p = doc.add_paragraph("Seu nome é ")
p.add_run(nome).bold = True #add_run: adiciona conteudo a um mesmo paragrafo (equivale ao append)
p.add_run('.')

doc.add_heading('A seguir outros detalhes:', level=1)

idade = input("Qual a sua idade? ")
doc.add_paragraph('Você tem {} anos de idade'.format(idade), style='List Bullet') #style='List Bullet': torna o parágrafo em formato de tópico

telefone = input("Qual é o seu telefone? ")
doc.add_paragraph('Seu telefone é o {}'.format(telefone), style='List Bullet')

endereco = input("Qual é o seu endereço? ")
doc.add_paragraph('Seu endereço é {}'.format(endereco), style='List Bullet')

foto = input("Digite o caminho (path) completo de um arquivo com sua foto? ")
doc.add_picture(foto, width=Inches(6), height=Inches(4))

# Aqui entrará o código de criação de tabela
tab = doc.add_table(rows=1, cols=2)
tab.style="Medium Grid 3 Accent 4" #Estilo da tabela
cels = tab.rows[0].cells #um array em python é como a tabela no word
cels[0].text = 'Prioridade'
cels[1].text = "Tarefas"

# Três tarefas importantes das pessoas em uma tabela
for numero in range(1, 4):
   tarefa = input("Qual é a tarefa de prioridade número-{}? ".format(numero))
   dados = tab.add_row().cells
   dados[0].text = str(numero)
   dados[1].text = tarefa


doc.save('meudoc.docx')
