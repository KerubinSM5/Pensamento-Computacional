from docx import Document
from utils import aleatorio

num_jogos = int(input('Quantas jogos quer jogar? '))
num_dezenas = int(input('Quantas dezenas quer jogar? '))

if num_jogos < 1:
    print("Numero invalido")
else:
    if num_dezenas < 6 or num_dezenas > 15:
        print("Não existem jogos com esse numero de dezenas")
    else:
        doc = Document()

        doc.add_heading('Bolão com {}'.format(num_jogos), 0)

        tab = doc.add_table(rows=1, cols=num_dezenas + 1)
        tab.style = "Medium List 1 Accent 3"
        cels = tab.rows[0].cells
        cels[0].text = 'Nº'
        for i in range(1, num_dezenas + 1):
            cels[i].text = f"Dez{i}"
        a = aleatorio()

        for rep in range(1, num_jogos + 1):
            dados = tab.add_row().cells
            dados[0].text = str(rep)
            seq = 1
            for dez in a.numeros(num_dezenas,1 ,60):
                dados[seq].text = str(dez)
                seq += 1

        doc.save(f'megasena{num_jogos}_{num_dezenas}.docx')
