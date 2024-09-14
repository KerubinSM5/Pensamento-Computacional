from docx import Document
from docx.shared import Inches

doc = Document()

doc.add_heading('Dados da disciplina', 0)

nome = input("Qual o nome da disciplina? ") #Nome da disciplina
p = doc.add_paragraph("Disciplina: ")
p.add_run(nome).bold = True #add_run: adiciona conteudo a um mesmo paragrafo (equivale ao append)

doc.add_heading('Detalhes:', level=1)

carga = input("Qual a carga horária? ")
p = doc.add_paragraph("Com carga horária de ")
p.add_run(carga).bold = True
p.add_run(' horas.')

tab = doc.add_table(rows=1, cols=2)
tab.style="Colorful Grid Accent 5"
cels = tab.rows[0].cells
cels[0].text = 'Número'
cels[0].width = 5
cels[1].text = "Assunto"

assunt = int(input("Quantos assuntos da assuntamento? "))
for numero in range(1, assunt+1):
   assunto = input("Qual é o assunto número-{}? ".format(numero))
   dados = tab.add_row().cells
   dados[0].text = str(numero)
   dados[0].width = 5
   dados[1].text = assunto

doc.save('meudocex.docx')