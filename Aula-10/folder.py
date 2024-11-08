from docx import Document
from util.db import SQL
from util import converter as cv




class Folder:
   def __init__(self):
       self.sql = SQL(esquema='bd_planejamento')
       self.arquivo = 'folder.docx'
       self.idt_projeto = 0


   def buscar_projeto(self, idt_projeto):
       cmd = """
          SELECT nme_projeto, dta_ini_projeto, dta_fim_projeto
            FROM tb_projeto
            WHERE idt_projeto = %s;"""
       projeto = self.sql.get_object(cmd, [idt_projeto])
       return projeto


   def buscar_colaboradores(self, idt_projeto):
       cmd = """
          SELECT nme_colaborador, eml_colaborador, nme_funcao, dta_ini_alocacao, dta_fim_alocacao
           FROM tb_colaborador JOIN ta_alocacao ON idt_colaborador = cod_colaborador
           JOIN tb_funcao ON idt_funcao = cod_funcao
           WHERE cod_projeto = %s
           ORDER BY nme_colaborador;"""
       lista = self.sql.get_list(cmd, [idt_projeto])
       return lista


   def gerar_folder(self):
       doc = Document()
       doc.add_heading('Projeto', 0)


       # Gerar a tabela de Projeto
       tb_proj = doc.add_table(rows=3, cols=2)
       tb_proj.style = "Medium List 1 Accent 2"
       projeto = self.buscar_projeto(self.idt_projeto)
       cels = tb_proj.rows[0].cells
       cels[0].text = 'Nome do Projeto'
       cels[1].text = projeto['nme_projeto']


       cels = tb_proj.rows[1].cells
       cels[0].text = 'Data de Início'
       cels[1].text = cv.mysql_to_bra(projeto['dta_ini_projeto'])


       cels = tb_proj.rows[2].cells
       cels[0].text = 'Data de Término'
       cels[1].text = cv.mysql_to_bra(projeto['dta_fim_projeto'])


       doc.add_heading('Colaboradores:', level=1)


       tb_col = doc.add_table(rows=1, cols=5)
       tb_col.style = "Medium List 2 Accent 2"
       cels = tb_col.rows[0].cells
       cels[0].text = 'Colaborador'
       cels[1].text = "e-mail"
       cels[2].text = "Função"
       cels[3].text = "Início"
       cels[4].text = "Término"


       lista = self.buscar_colaboradores(self.idt_projeto)
       for obj in lista:
           dados = tb_col.add_row().cells
           dados[0].text = obj['nme_colaborador']
           dados[1].text = obj['eml_colaborador']
           dados[2].text = obj['nme_funcao']
           dados[3].text = cv.mysql_to_bra(obj['dta_ini_alocacao'])
           dados[4].text = cv.mysql_to_bra(obj['dta_fim_alocacao'])


       doc.save(self.arquivo)




if __name__ == '__main__':
   folder = Folder()
   folder.idt_projeto = 2
   folder.gerar_folder()